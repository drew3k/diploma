from __future__ import annotations
from pathlib import Path
from typing import List
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import qn
from docx import Document
from .utils import Span, smart_mask, mask_person, DEFAULT_REGEX

import regex as re  # unicode-friendly regex

from natasha import NamesExtractor, MorphVocab
import pymorphy2

_morph_vocab = MorphVocab()
names_extractor = NamesExtractor(_morph_vocab)
morph = pymorphy2.MorphAnalyzer()


OOXML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

# --- служебные классы символов ---
CTRL0 = "\u200b\u200c\u200d\ufeff\u00ad"  # ZWSP/ZWNJ/ZWJ/BOM/soft hyphen
NBSP = "\u00a0"
SOFT_SEP_RX = re.compile(rf"[\s{NBSP}\t{re.escape(CTRL0)}]+")
SOFT_SEP_OR_EMPTY_RX = re.compile(rf"[\s{NBSP}\t{re.escape(CTRL0)}]*")

# слово целиком кириллицей (включаем дефисные части), без цифр/латиницы
WORD_RX = re.compile(r"[А-Яа-яЁё]+(?:-[А-Яа-яЁё]+)*")

# Отчества по продуктивным суффиксам
PATR_RX = re.compile(r".+?(?:ович|евич|ич|овна|евна|ична)$", re.IGNORECASE)

# Инициалы, допускаем «мягкие» разделители между точками
SOFT = rf"[\s{NBSP}{re.escape(CTRL0)}]"
INITIALS_RX = re.compile(rf"[А-ЯЁ]\.\s*{SOFT}?\s*[А-ЯЁ]\.", re.UNICODE)

SURNAME_RX = re.compile(r"[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?")


# --- вспомогательные функции DOCX ----
def _strip_comments_and_tracked_changes(doc: Document) -> None:
    root = doc._element.getroottree().getroot()
    for tag in (
        ".//w:commentRangeStart",
        ".//w:commentRangeEnd",
        ".//w:commentReference",
    ):
        for el in root.findall(tag, namespaces=OOXML_NS):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
    for el in root.findall(".//w:del", namespaces=OOXML_NS):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _span_rx(text: str) -> re.Pattern:
    esc = re.escape(text)
    esc = esc.replace(r"\ ", rf"[\s{NBSP}\t{re.escape(CTRL0)}]+")
    return re.compile(esc)


def _replace_by_spans_in_runs(runs, spans: List[Span], policy: str) -> int:
    count = 0
    for run in runs:
        t = run.text or ""
        changed = False
        for s in spans:
            if not s.text:
                continue
            rx = _span_rx(s.text)
            repl = smart_mask(s.label, s.text) if policy == "mask" else ""
            new_t, n = rx.subn(repl, t)
            if n:
                t = new_t
                changed = True
                count += n
        if changed:
            run.text = t
    return count


# --------- НОВОЕ: извлечение ФИО не по «кап-словам», а по морфологии ---------
def _has_name_gram(token: str) -> bool:
    tok = re.sub(f"[{re.escape(CTRL0)}]", "", token)
    for p in morph.parse(tok):
        t = str(p.tag)
        if ("Name" in t) or ("Surn" in t) or ("Patr" in t):
            return True
    return False


def _extract_person_spans_morph(text: str) -> List[Span]:
    spans: List[Span] = []

    # 1) Natasha: устойчивые антропонимы (Имя+Фамилия, ФИО и т.п.)
    for m in names_extractor(text):
        chunk = text[m.start : m.stop]
        parts = [t for t in re.split(SOFT_SEP_RX, chunk) if t]
        # простая валидация: минимум два токена с граммемами Name/Surn/Patr или валидное отчество
        votes = sum(
            1 for t in parts if (_has_name_gram(t) or bool(PATR_RX.fullmatch(t)))
        )
        if votes >= 2:
            spans.append(Span(chunk, m.start, m.stop, "PERSON"))

    # 2) Фамилия + инициалы / инициалы + фамилия
    pattern1 = re.compile(
        rf"({SURNAME_RX.pattern}){SOFT_SEP_OR_EMPTY_RX.pattern}({INITIALS_RX.pattern})"
    )
    pattern2 = re.compile(
        rf"({INITIALS_RX.pattern}){SOFT_SEP_OR_EMPTY_RX.pattern}({SURNAME_RX.pattern})"
    )
    for pat in (pattern1, pattern2):
        for m in pat.finditer(text):
            spans.append(Span(m.group(0), m.start(), m.end(), "PERSON"))

    # dedup вложенных
    spans.sort(key=lambda s: (s.start, s.end))
    ded: List[Span] = []
    for s in spans:
        if ded and s.start >= ded[-1].start and s.end <= ded[-1].end:
            continue
        ded.append(s)
    return ded


def _mask_text_by_spans(text: str, spans: List[Span]) -> tuple[str, int]:
    """Заменяем в исходной строке только указанные диапазоны PERSON → mask_person()."""
    if not spans:
        return text, 0
    spans = sorted(spans, key=lambda s: s.start)
    out = []
    pos = 0
    count = 0
    for s in spans:
        if s.start < pos:
            continue
        out.append(text[pos : s.start])
        out.append(mask_person(text[s.start : s.end]))
        pos = s.end
        count += 1
    out.append(text[pos:])
    return "".join(out), count


def _fallback_mask_paragraph_person_only(p) -> int:
    """
    Новый fallback: в рамках абзаца извлекаем PERSON по Natasha+pymorphy2
    (плюс фамилия+инициалы), маскируем только их. Обычный текст не трогаем.
    """
    if not p.runs:
        return 0
    original = "".join(r.text or "" for r in p.runs)
    spans = _extract_person_spans_morph(original)
    if not spans:
        return 0
    masked, n = _mask_text_by_spans(original, spans)
    if n == 0 or masked == original:
        return 0
    p.runs[0].text = masked
    for r in p.runs[1:]:
        r.text = ""
    return n


# --------- Применение к разделам документа ---------
def _replace_in_paragraph_obj(p, spans: List[Span], policy: str) -> int:
    total = 0
    total += _replace_by_spans_in_runs(p.runs, spans, policy)
    try:
        hls = getattr(p._element, "hyperlink_lst", [])
        for hl in hls:
            runs_in_hl = getattr(hl, "r_lst", [])
            total += _replace_by_spans_in_runs(runs_in_hl, spans, policy)
    except Exception:
        pass
    if policy == "mask":
        total += _fallback_mask_paragraph_person_only(p)
    return total


def _replace_in_paragraphs(doc: Document, spans: List[Span], policy: str) -> int:
    total = 0
    for p in doc.paragraphs:
        total += _replace_in_paragraph_obj(p, spans, policy)
    return total


def _replace_in_headers_footers(doc: Document, spans: List[Span], policy: str) -> int:
    total = 0
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                total += _replace_in_paragraph_obj(p, spans, policy)
        if section.footer:
            for p in section.footer.paragraphs:
                total += _replace_in_paragraph_obj(p, spans, policy)
    return total


def _replace_in_tables(doc: Document, spans: List[Span], policy: str) -> int:
    total = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += _replace_in_paragraph_obj(p, spans, policy)
    return total


def _sanitize_hyperlinks_with_emails(doc: Document, policy) -> int:
    """
    Маскирует e-mail в:
      - target_ref у внешних гиперссылок (mailto:...),
      - видимом тексте внутри <w:hyperlink>,
      - полях HYPERLINK (w:instrText), содержащих mailto:.

    policy.email.mode: "mask" | "remove"  (если "remove": снимаем кликабельность)
    Возвращает счётчик модифицированных мест.
    """
    # --- настройки/регексы из твоего utils.py ---
    EMAIL_RX = DEFAULT_REGEX["EMAIL_ADDRESS"]
    mode = getattr(getattr(policy, "email", object()), "mode", "mask")

    count = 0

    # 1) Исправляем целевые URL у внешних ссылок reltype=HYPERLINK (mailto:)
    rels = getattr(doc.part, "rels", {})
    for rel_id, rel in list(rels.items()):
        try:
            if (
                rel.reltype == RT.HYPERLINK
                and isinstance(rel.target_ref, str)
                and rel.target_ref.lower().startswith("mailto:")
            ):
                addr = rel.target_ref[7:]
                if EMAIL_RX.search(addr):
                    if mode == "mask":
                        masked = smart_mask("EMAIL_ADDRESS", addr)
                        new_target = "mailto:" + masked
                        if new_target != rel.target_ref:
                            rel._target = new_target  # python-docx внутренний setter
                            count += 1
                    else:  # remove -> делаем ссылку некликабельной, текст обработаем ниже
                        # Отрываем relationship, чтобы ссылка не была кликабельной
                        rels.pop(rel_id, None)
                        count += 1
        except Exception:
            # защитно: не падаем на экзотических rel
            continue

    # 2) Маскируем видимый текст внутри <w:hyperlink> даже если он разбит на несколько run
    root = doc.part.element  # это <w:document>
    for hl in root.findall(".//w:hyperlink", namespaces=OOXML_NS):
        # Собираем все текстовые узлы под гиперссылкой в порядке следования
        t_nodes = hl.findall(".//w:r//w:t", namespaces=OOXML_NS)
        if not t_nodes:
            continue
        visible = "".join([(t.text or "") for t in t_nodes])
        if not visible:
            continue

        # Есть ли e-mail в видимом тексте?
        if not EMAIL_RX.search(visible):
            continue

        # Маскируем все вхождения e-mail в склейке
        def _mask_all(text: str) -> str:
            out, i = [], 0
            for m in EMAIL_RX.finditer(text):
                out.append(text[i : m.start()])
                out.append(smart_mask("EMAIL_ADDRESS", m.group(0)))
                i = m.end()
            out.append(text[i:])
            return "".join(out)

        masked = _mask_all(visible)
        if masked != visible:
            # Записываем обратно: в первый <w:t> кладём всю строку, остальные чистим
            t_nodes[0].text = masked
            for t in t_nodes[1:]:
                t.text = ""
            count += 1

        # Если режим remove — дополнительно снимаем r:id у самой гиперссылки
        if mode != "mask":
            r_id = hl.get(qn("r:id"))
            if r_id is not None:
                hl.attrib.pop(qn("r:id"), None)
                count += 1

    # 3) Поля HYPERLINK с mailto: внутри w:instrText
    # Пример содержимого: ' HYPERLINK "mailto:john.doe@mail.ru" '
    for instr in root.findall(".//w:instrText", namespaces=OOXML_NS):
        text = instr.text or ""
        if "HYPERLINK" not in text or "mailto:" not in text:
            continue

        # Ищем внутри кавычек адрес и маскируем
        # Берём самый частый вариант: ... "mailto:...@..." ...
        def _replace_mailto(mo):
            whole = mo.group(0)  # '"mailto:addr"'
            addr = mo.group("addr")  # 'addr'
            if not EMAIL_RX.search(addr):
                return whole
            if mode == "mask":
                return f'"mailto:{smart_mask("EMAIL_ADDRESS", addr)}"'
            else:
                # В remove режиме оставим поле, но с некликабельным pseudo-URL,
                # либо можно вычистить почту целиком. Выберем безопасный вариант — маска.
                return f'"mailto:{smart_mask("EMAIL_ADDRESS", addr)}"'

        new_text = re.sub(
            r"\"mailto:(?P<addr>[^\"]+?)\"",
            _replace_mailto,
            text,
        )
        if new_text != text:
            instr.text = new_text
            count += 1

    return count


def _clear_core_properties(doc: Document) -> None:
    cp = doc.core_properties
    for attr in (
        "author",
        "last_modified_by",
        "category",
        "comments",
        "keywords",
        "subject",
        "title",
    ):
        try:
            setattr(cp, attr, "")
        except Exception:
            pass


def cleanse_docx(inp: Path, out: Path, spans: List[Span], policy: str = "mask") -> int:
    doc = Document(str(inp))

    _strip_comments_and_tracked_changes(doc)

    replaced = 0
    replaced += _replace_in_headers_footers(doc, spans, policy)
    replaced += _replace_in_paragraphs(doc, spans, policy)
    replaced += _replace_in_tables(doc, spans, policy)
    replaced += _sanitize_hyperlinks_with_emails(doc, policy)

    _clear_core_properties(doc)
    doc.save(str(out))
    return replaced
