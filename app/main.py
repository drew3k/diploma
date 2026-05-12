from __future__ import annotations

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from starlette.responses import StreamingResponse
from typing import List
from pathlib import Path
from io import BytesIO
from secrets import token_urlsafe
from urllib.parse import quote
import unicodedata
from app.pipeline.feedback import log_candidates
from app.models import ProcessResult
from app.settings import settings
from app.storage import save_upload, out_path_for, public_url, OUT
from app.pipeline.detect import detect_spans
from app.pipeline.utils import DEFAULT_REGEX
from app.pipeline.redact_pdf import redact_pdf
from app.pipeline.cleanse_docx import cleanse_docx

app = FastAPI(title="PD Redactor Service", version="0.2.2")

# --------------------------- Стартовая страница ----------------------------
INDEX = (Path(__file__).parent / "static" / "index.html").read_text(encoding="utf-8")


@app.get("/", response_class=HTMLResponse)
def index():
    return INDEX


@app.get("/api/health")
def health():
    return {"status": "ok"}


# --------------------------- Классический API (сохранение на диск) --------
@app.post("/api/process", response_model=List[ProcessResult])
async def process_files(
    files: List[UploadFile] = File(...),
    policy: str = Form("mask"),
    languages: str = Form("ru,en"),
    types: str | None = Form(None),
):
    if policy not in ("mask", "remove"):
        raise HTTPException(400, "policy must be 'mask' or 'remove'")

    results: List[ProcessResult] = []

    for uf in files:
        content = await uf.read()
        if len(content) > settings.max_file_mb * 1024 * 1024:
            raise HTTPException(413, f"{uf.filename}: file too large")

        inp = save_upload(uf.filename, content)
        suffix = Path(uf.filename).suffix.lower()

        if suffix == ".pdf":
            import fitz

            with fitz.open(inp) as doc:
                text_for_detection = "\n".join(page.get_text() for page in doc)
        elif suffix == ".docx":
            from docx import Document

            d = Document(str(inp))
            # Base text from paragraphs
            text_for_detection = "\n".join(p.text for p in d.paragraphs)

            # Also include emails from hyperlink targets (e.g., mailto: links)
            try:
                EMAIL_RX = DEFAULT_REGEX.get("EMAIL_ADDRESS")
                if EMAIL_RX is not None:

                    def _emails_from_part(part) -> list[str]:
                        out: list[str] = []
                        try:
                            for rel in part.rels.values():
                                # Word hyperlink relationships
                                if (
                                    rel.reltype
                                    == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
                                ):
                                    target = getattr(rel, "target_ref", "") or ""
                                    m = EMAIL_RX.search(target)
                                    if m:
                                        out.append(m.group(0))
                        except Exception:
                            pass
                        return out

                    emails: set[str] = set(_emails_from_part(d.part))
                    # headers/footers have separate parts and rels
                    for section in d.sections:
                        if getattr(section, "header", None):
                            emails.update(_emails_from_part(section.header.part))
                        if getattr(section, "footer", None):
                            emails.update(_emails_from_part(section.footer.part))
                    if emails:
                        text_for_detection += "\n" + "\n".join(sorted(emails))
            except Exception:
                # Best-effort: ignore hyperlink extraction failures
                pass
        else:
            raise HTTPException(415, f"Unsupported type: {suffix}")

        spans = detect_spans(
            text_for_detection,
            languages.split(","),
            set(types.split(",")) if types else None,
        )

        log_candidates(text_for_detection, spans, source="api")

        if suffix == ".pdf":
            out = out_path_for(inp, ".pdf")
            found = redact_pdf(inp, out, spans, policy)
            results.append(
                ProcessResult(
                    input_name=uf.filename,
                    output_name=out.name,
                    output_url=public_url(out),
                    found=found,
                    filetype="pdf",
                )
            )
        else:
            out = out_path_for(inp, ".docx")
            found = cleanse_docx(inp, out, spans, policy)
            results.append(
                ProcessResult(
                    input_name=uf.filename,
                    output_name=out.name,
                    output_url=public_url(out),
                    found=found,
                    filetype="docx",
                )
            )

    return results


@app.get("/api/file/{name}")
def download_saved(name: str):
    p = OUT / name
    if not p.exists():
        raise HTTPException(404)
    return FileResponse(p)


# --------------------------- Обработка в памяти + красивый скачиватор -----

# Простая in-memory кэш-таблица: token -> {"bytes":..., "media":..., "filename":...}
DOWNLOAD_CACHE: dict[str, dict[str, bytes | str]] = {}


def _detect_text_from_pdf_bytes(data: bytes) -> str:
    import fitz

    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _redact_pdf_in_memory(data: bytes, spans, policy: str) -> bytes:
    import fitz
    from app.pipeline.utils import smart_mask

    BLACK = (0, 0, 0)
    WHITE = (1, 1, 1)

    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            for s in spans:
                txt = (s.text or "").strip()
                if not txt:
                    continue
                rects = page.search_for(txt, quads=False)
                if not rects:
                    continue
                overlay_text = smart_mask(s.label, s.text) if policy == "mask" else None
                fill_color = WHITE if policy == "mask" else BLACK
                for r in rects:
                    page.add_redact_annot(r, fill=fill_color, text=overlay_text)
            page.apply_redactions(images=(policy == "remove"))

        try:
            doc.set_metadata({})
        except Exception:
            pass

        return doc.tobytes(garbage=4, deflate=True, clean=True)


def _detect_text_from_docx_bytes(data: bytes) -> str:
    from docx import Document

    d = Document(BytesIO(data))
    # Base text from paragraphs
    text = "\n".join(p.text for p in d.paragraphs)

    # Also include emails from hyperlink targets (e.g., mailto: links)
    try:
        EMAIL_RX = DEFAULT_REGEX.get("EMAIL_ADDRESS")
        if EMAIL_RX is not None:

            def _emails_from_part(part) -> list[str]:
                out: list[str] = []
                try:
                    for rel in part.rels.values():
                        if (
                            rel.reltype
                            == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
                        ):
                            target = getattr(rel, "target_ref", "") or ""
                            m = EMAIL_RX.search(target)
                            if m:
                                out.append(m.group(0))
                except Exception:
                    pass
                return out

            emails: set[str] = set(_emails_from_part(d.part))
            for section in d.sections:
                if getattr(section, "header", None):
                    emails.update(_emails_from_part(section.header.part))
                if getattr(section, "footer", None):
                    emails.update(_emails_from_part(section.footer.part))
            if emails:
                text += "\n" + "\n".join(sorted(emails))
    except Exception:
        pass

    return text


def _cleanse_docx_in_memory(data: bytes, spans, policy: str) -> bytes:
    from docx import Document
    from app.pipeline.cleanse_docx import (
        _replace_in_headers_footers,
        _replace_in_paragraphs,
        _replace_in_tables,
        _clear_core_properties,
        _strip_comments_and_tracked_changes,
        _sanitize_hyperlinks_with_emails,
    )

    doc = Document(BytesIO(data))
    _strip_comments_and_tracked_changes(doc)
    _replace_in_headers_footers(doc, spans, policy)
    _replace_in_paragraphs(doc, spans, policy)
    _replace_in_tables(doc, spans, policy)
    _sanitize_hyperlinks_with_emails(doc, policy)
    _clear_core_properties(doc)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/web/submit")
async def web_submit(
    files: List[UploadFile] = File(...),
    policy: str = Form("mask"),
    languages: str = Form("ru,en"),
    types: str | None = Form(None),
):
    if policy not in ("mask", "remove"):
        raise HTTPException(400, "policy must be 'mask' or 'remove'")

    processed_items: List[dict[str, str]] = []

    for file in files:
        content = await file.read()
        suffix = Path(file.filename).suffix.lower()

        if suffix == ".pdf":
            text_for_detection = _detect_text_from_pdf_bytes(content)
        elif suffix == ".docx":
            text_for_detection = _detect_text_from_docx_bytes(content)
        else:
            raise HTTPException(415, f"Unsupported type: {suffix}")

        spans = detect_spans(
            text_for_detection,
            languages.split(","),
            set(types.split(",")) if types else None,
        )

        log_candidates(text_for_detection, spans, source="web")

        if suffix == ".pdf":
            processed = _redact_pdf_in_memory(content, spans, policy)
            media = "application/pdf"
        else:
            processed = _cleanse_docx_in_memory(content, spans, policy)
            media = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        token = token_urlsafe(16)
        download_name = f"{Path(file.filename).stem}_redacted{suffix}"
        DOWNLOAD_CACHE[token] = {
            "bytes": processed,
            "media": media,
            "filename": download_name,
        }
        processed_items.append({"token": token, "filename": download_name})

    if len(processed_items) == 1:
        return RedirectResponse(url=f"/download/{processed_items[0]['token']}", status_code=303)

    links = "".join(
    f"""
    <a class="file-link" href="/api/download/{item["token"]}">
        <span class="file-icon">📄</span>
        <span class="file-info">
            <strong>{item["filename"]}</strong>
            <small>Скачать обезличенный документ</small>
        </span>
        <span class="download-icon">⬇</span>
    </a>
    """
    for item in processed_items
    )
    html = f"""
<!doctype html>
<html lang="ru">
<head>
    <meta charset="utf-8"/>
    <meta name="viewport" content="width=device-width, initial-scale=1"/>
    <title>Файлы готовы — PD Redactor</title>

    <style>
        :root {{
            --bg-1: #080b18;
            --bg-2: #101936;
            --text: #f8fafc;
            --muted: #a9b4c7;
            --blue: #6ee7ff;
            --violet: #8b5cf6;
            --pink: #ec4899;
            --green: #34d399;
            --card-border: rgba(255, 255, 255, 0.18);
            --shadow: 0 30px 80px rgba(0, 0, 0, 0.35);
        }}

        * {{
            box-sizing: border-box;
        }}

        body {{
            margin: 0;
            min-height: 100vh;
            color: var(--text);
            font-family:
                Inter,
                ui-sans-serif,
                system-ui,
                -apple-system,
                BlinkMacSystemFont,
                "Segoe UI",
                Roboto,
                Arial,
                sans-serif;
            background:
                radial-gradient(circle at 15% 15%, rgba(110, 231, 255, 0.22), transparent 28%),
                radial-gradient(circle at 80% 5%, rgba(139, 92, 246, 0.28), transparent 30%),
                radial-gradient(circle at 80% 85%, rgba(236, 72, 153, 0.18), transparent 30%),
                linear-gradient(135deg, var(--bg-1), var(--bg-2));
            overflow-x: hidden;
        }}

        body::before {{
            content: "";
            position: fixed;
            inset: 0;
            pointer-events: none;
            background-image:
                linear-gradient(rgba(255, 255, 255, 0.035) 1px, transparent 1px),
                linear-gradient(90deg, rgba(255, 255, 255, 0.035) 1px, transparent 1px);
            background-size: 42px 42px;
            mask-image: linear-gradient(to bottom, rgba(0, 0, 0, 0.9), transparent);
        }}

        .orb {{
            position: fixed;
            border-radius: 999px;
            filter: blur(12px);
            opacity: 0.6;
            pointer-events: none;
            animation: float 8s ease-in-out infinite;
        }}

        .orb.one {{
            width: 210px;
            height: 210px;
            left: -60px;
            top: 190px;
            background: rgba(110, 231, 255, 0.28);
        }}

        .orb.two {{
            width: 260px;
            height: 260px;
            right: -90px;
            top: 130px;
            background: rgba(139, 92, 246, 0.32);
            animation-delay: -2s;
        }}

        .orb.three {{
            width: 180px;
            height: 180px;
            right: 18%;
            bottom: -70px;
            background: rgba(236, 72, 153, 0.24);
            animation-delay: -4s;
        }}

        @keyframes float {{
            0%, 100% {{
                transform: translate3d(0, 0, 0);
            }}

            50% {{
                transform: translate3d(0, -24px, 0);
            }}
        }}

        .page {{
            position: relative;
            z-index: 1;
            width: min(1180px, calc(100% - 32px));
            margin: 0 auto;
            padding: 34px 0 56px;
        }}

        .topbar {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 18px;
            margin-bottom: 58px;
        }}

        .brand {{
            display: flex;
            align-items: center;
            gap: 12px;
            font-weight: 800;
            letter-spacing: 0.2px;
        }}

        .brand-icon {{
            width: 44px;
            height: 44px;
            display: grid;
            place-items: center;
            border-radius: 16px;
            background:
                linear-gradient(135deg, rgba(110, 231, 255, 0.95), rgba(139, 92, 246, 0.95));
            box-shadow: 0 14px 38px rgba(110, 231, 255, 0.2);
        }}

        .status-pill {{
            display: inline-flex;
            align-items: center;
            gap: 9px;
            padding: 10px 14px;
            border: 1px solid rgba(255, 255, 255, 0.14);
            border-radius: 999px;
            color: rgba(255, 255, 255, 0.78);
            background: rgba(255, 255, 255, 0.07);
            backdrop-filter: blur(18px);
            font-size: 14px;
        }}

        .pulse {{
            width: 9px;
            height: 9px;
            border-radius: 50%;
            background: var(--green);
            box-shadow: 0 0 0 6px rgba(52, 211, 153, 0.14);
        }}

        .result-wrap {{
            display: grid;
            place-items: center;
            min-height: calc(100vh - 190px);
        }}

        .panel {{
            width: min(820px, 100%);
            position: relative;
            padding: 1px;
            border-radius: 34px;
            background:
                linear-gradient(145deg, rgba(110, 231, 255, 0.55), rgba(139, 92, 246, 0.2), rgba(236, 72, 153, 0.45));
            box-shadow: var(--shadow);
        }}

        .panel-inner {{
            position: relative;
            padding: 42px;
            border-radius: 33px;
            border: 1px solid var(--card-border);
            background:
                linear-gradient(180deg, rgba(255, 255, 255, 0.13), rgba(255, 255, 255, 0.07));
            backdrop-filter: blur(24px);
            overflow: hidden;
            text-align: center;
        }}

        .panel-inner::before {{
            content: "";
            position: absolute;
            inset: -100px -120px auto auto;
            width: 280px;
            height: 280px;
            background: radial-gradient(circle, rgba(110, 231, 255, 0.18), transparent 68%);
            pointer-events: none;
        }}

        .success-icon {{
            position: relative;
            width: 92px;
            height: 92px;
            display: grid;
            place-items: center;
            margin: 0 auto 24px;
            border-radius: 30px;
            background:
                linear-gradient(135deg, rgba(52, 211, 153, 0.95), rgba(110, 231, 255, 0.95));
            box-shadow: 0 22px 55px rgba(52, 211, 153, 0.2);
        }}

        h1 {{
            position: relative;
            margin: 0;
            font-size: clamp(36px, 4.4vw, 54px);
            line-height: 1.05;
            letter-spacing: -1.6px;
        }}

        .gradient-text {{
            color: #67e8f9;
            text-shadow: 0 0 22px rgba(103, 232, 249, 0.45);
        }}

        .subtitle {{
            position: relative;
            margin: 18px auto 0;
            max-width: 620px;
            color: var(--muted);
            font-size: 17px;
            line-height: 1.65;
        }}

        .files-list {{
            position: relative;
            display: grid;
            gap: 12px;
            margin: 30px auto 0;
            width: min(560px, 100%);
        }}

        .file-link {{
            display: grid;
            grid-template-columns: 44px 1fr 34px;
            align-items: center;
            gap: 14px;
            min-height: 72px;
            padding: 13px 15px;
            border-radius: 20px;
            color: var(--text);
            text-decoration: none;
            text-align: left;
            background:
                linear-gradient(135deg, rgba(110, 231, 255, 0.09), rgba(236, 72, 153, 0.08)),
                rgba(255, 255, 255, 0.08);
            border: 1px solid rgba(255, 255, 255, 0.14);
            transition: 0.22s ease;
        }}

        .file-link:hover {{
            transform: translateY(-2px);
            border-color: rgba(110, 231, 255, 0.75);
            box-shadow:
                0 18px 46px rgba(139, 92, 246, 0.22),
                0 0 0 4px rgba(110, 231, 255, 0.08);
        }}

        .file-icon {{
            width: 44px;
            height: 44px;
            display: grid;
            place-items: center;
            border-radius: 16px;
            background:
                linear-gradient(135deg, rgba(110, 231, 255, 0.22), rgba(139, 92, 246, 0.22));
            border: 1px solid rgba(255, 255, 255, 0.12);
        }}

        .file-info {{
            min-width: 0;
            display: grid;
            gap: 4px;
        }}

        .file-info strong {{
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
            font-size: 15px;
        }}

        .file-info small {{
            color: var(--muted);
            font-size: 13px;
        }}

        .download-icon {{
            width: 34px;
            height: 34px;
            display: grid;
            place-items: center;
            border-radius: 13px;
            color: #06101f;
            background: linear-gradient(135deg, #67e8f9, #a78bfa 55%, #f0abfc);
            font-weight: 900;
        }}

        .actions {{
            position: relative;
            margin-top: 30px;
        }}

        .btn-secondary {{
            min-height: 54px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            padding: 0 22px;
            border-radius: 18px;
            border: 1px solid rgba(255, 255, 255, 0.16);
            color: var(--text);
            background: rgba(255, 255, 255, 0.08);
            text-decoration: none;
            font-weight: 800;
            transition: 0.2s ease;
        }}

        .btn-secondary:hover {{
            transform: translateY(-2px);
            background: rgba(255, 255, 255, 0.13);
        }}

        @media (max-width: 720px) {{
            .topbar {{
                align-items: flex-start;
                flex-direction: column;
                margin-bottom: 36px;
            }}

            .panel-inner {{
                padding: 28px 20px;
            }}

            .file-link {{
                grid-template-columns: 40px 1fr;
            }}

            .download-icon {{
                display: none;
            }}

            .btn-secondary {{
                width: 100%;
            }}
        }}
    </style>
</head>

<body>
    <div class="orb one"></div>
    <div class="orb two"></div>
    <div class="orb three"></div>

    <main class="page">
        <header class="topbar">
            <div class="brand">
                <div class="brand-icon">
                    <svg width="24" height="24" viewBox="0 0 24 24" fill="none">
                        <path d="M12 3L20 7V12C20 17 16.5 20.5 12 21C7.5 20.5 4 17 4 12V7L12 3Z" stroke="white" stroke-width="2" stroke-linejoin="round"/>
                        <path d="M9 12L11 14L15.5 9.5" stroke="white" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
                    </svg>
                </div>
                <div>
                    <div>PD Redactor</div>
                    <small style="color: var(--muted); font-weight: 600;">Secure document anonymization</small>
                </div>
            </div>

            <div class="status-pill">
                <span class="pulse"></span>
                Обработка завершена
            </div>
        </header>

        <section class="result-wrap">
            <div class="panel">
                <div class="panel-inner">
                    <div class="success-icon">
                        <svg width="44" height="44" viewBox="0 0 24 24" fill="none">
                            <path d="M20 6L9 17L4 12" stroke="white" stroke-width="2.7" stroke-linecap="round" stroke-linejoin="round"/>
                        </svg>
                    </div>

                    <h1>
                        Файлы <span class="gradient-text">готовы</span>
                    </h1>

                    <p class="subtitle">
                        Все загруженные документы успешно обработаны. Скачайте обезличенные версии файлов по ссылкам ниже.
                    </p>

                    <div class="files-list">
                        {links}
                    </div>

                    <div class="actions">
                        <a class="btn-secondary" href="/">Обработать ещё документы</a>
                    </div>
                </div>
            </div>
        </section>
    </main>
</body>
</html>
"""
    return HTMLResponse(html)


@app.get("/download/{token}", response_class=HTMLResponse)
def download_page(token: str):
    item = DOWNLOAD_CACHE.get(token)
    if not item:
        html = """
<!doctype html>
<html lang="ru">
<head>
    <meta charset="utf-8"/>
    <meta name="viewport" content="width=device-width, initial-scale=1"/>
    <title>Ссылка не найдена</title>
    <style>
        body {
            margin: 0;
            min-height: 100vh;
            display: grid;
            place-items: center;
            color: #f8fafc;
            font-family: Inter, ui-sans-serif, system-ui, "Segoe UI", Roboto, Arial;
            background:
                radial-gradient(circle at 15% 15%, rgba(110, 231, 255, 0.22), transparent 28%),
                radial-gradient(circle at 80% 5%, rgba(139, 92, 246, 0.28), transparent 30%),
                linear-gradient(135deg, #080b18, #101936);
        }
        .card {
            width: min(620px, calc(100% - 32px));
            padding: 36px;
            border-radius: 30px;
            text-align: center;
            background: rgba(255,255,255,.1);
            border: 1px solid rgba(255,255,255,.18);
            backdrop-filter: blur(24px);
            box-shadow: 0 30px 80px rgba(0,0,0,.35);
        }
        h1 { margin: 0 0 10px; font-size: 36px; }
        p { margin: 0 0 26px; color: #a9b4c7; line-height: 1.6; }
        a {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-height: 52px;
            padding: 0 22px;
            border-radius: 17px;
            color: #06101f;
            font-weight: 900;
            text-decoration: none;
            background: linear-gradient(135deg, #67e8f9, #a78bfa 55%, #f0abfc);
        }
    </style>
</head>
<body>
    <div class="card">
        <h1>Ссылка устарела</h1>
        <p>Файл не найден в кэше сервера. Загрузите документ повторно и запустите обработку ещё раз.</p>
        <a href="/">Вернуться на главную</a>
    </div>
</body>
</html>
"""
        return HTMLResponse(html, status_code=404)

    filename = item["filename"]

    html = f"""
<!doctype html>
<html lang="ru">
<head>
    <meta charset="utf-8"/>
    <meta name="viewport" content="width=device-width, initial-scale=1"/>
    <title>Файл готов — PD Redactor</title>

    <style>
        :root {{
            --bg-1: #080b18;
            --bg-2: #101936;
            --text: #f8fafc;
            --muted: #a9b4c7;
            --blue: #6ee7ff;
            --violet: #8b5cf6;
            --pink: #ec4899;
            --green: #34d399;
            --card-border: rgba(255, 255, 255, 0.18);
            --shadow: 0 30px 80px rgba(0, 0, 0, 0.35);
        }}

        * {{
            box-sizing: border-box;
        }}

        body {{
            margin: 0;
            min-height: 100vh;
            color: var(--text);
            font-family:
                Inter,
                ui-sans-serif,
                system-ui,
                -apple-system,
                BlinkMacSystemFont,
                "Segoe UI",
                Roboto,
                Arial,
                sans-serif;
            background:
                radial-gradient(circle at 15% 15%, rgba(110, 231, 255, 0.22), transparent 28%),
                radial-gradient(circle at 80% 5%, rgba(139, 92, 246, 0.28), transparent 30%),
                radial-gradient(circle at 80% 85%, rgba(236, 72, 153, 0.18), transparent 30%),
                linear-gradient(135deg, var(--bg-1), var(--bg-2));
            overflow-x: hidden;
        }}

        body::before {{
            content: "";
            position: fixed;
            inset: 0;
            pointer-events: none;
            background-image:
                linear-gradient(rgba(255, 255, 255, 0.035) 1px, transparent 1px),
                linear-gradient(90deg, rgba(255, 255, 255, 0.035) 1px, transparent 1px);
            background-size: 42px 42px;
            mask-image: linear-gradient(to bottom, rgba(0, 0, 0, 0.9), transparent);
        }}

        .orb {{
            position: fixed;
            border-radius: 999px;
            filter: blur(12px);
            opacity: 0.6;
            pointer-events: none;
            animation: float 8s ease-in-out infinite;
        }}

        .orb.one {{
            width: 210px;
            height: 210px;
            left: -60px;
            top: 190px;
            background: rgba(110, 231, 255, 0.28);
        }}

        .orb.two {{
            width: 260px;
            height: 260px;
            right: -90px;
            top: 130px;
            background: rgba(139, 92, 246, 0.32);
            animation-delay: -2s;
        }}

        .orb.three {{
            width: 180px;
            height: 180px;
            right: 18%;
            bottom: -70px;
            background: rgba(236, 72, 153, 0.24);
            animation-delay: -4s;
        }}

        @keyframes float {{
            0%, 100% {{
                transform: translate3d(0, 0, 0);
            }}
            50% {{
                transform: translate3d(0, -24px, 0);
            }}
        }}

        .page {{
            position: relative;
            z-index: 1;
            width: min(1180px, calc(100% - 32px));
            margin: 0 auto;
            padding: 34px 0 56px;
        }}

        .topbar {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 18px;
            margin-bottom: 70px;
        }}

        .brand {{
            display: flex;
            align-items: center;
            gap: 12px;
            font-weight: 800;
            letter-spacing: 0.2px;
        }}

        .brand-icon {{
            width: 44px;
            height: 44px;
            display: grid;
            place-items: center;
            border-radius: 16px;
            background:
                linear-gradient(135deg, rgba(110, 231, 255, 0.95), rgba(139, 92, 246, 0.95));
            box-shadow: 0 14px 38px rgba(110, 231, 255, 0.2);
        }}

        .status-pill {{
            display: inline-flex;
            align-items: center;
            gap: 9px;
            padding: 10px 14px;
            border: 1px solid rgba(255, 255, 255, 0.14);
            border-radius: 999px;
            color: rgba(255, 255, 255, 0.78);
            background: rgba(255, 255, 255, 0.07);
            backdrop-filter: blur(18px);
            font-size: 14px;
        }}

        .pulse {{
            width: 9px;
            height: 9px;
            border-radius: 50%;
            background: var(--green);
            box-shadow: 0 0 0 6px rgba(52, 211, 153, 0.14);
        }}

        .result-wrap {{
            display: grid;
            place-items: center;
            min-height: calc(100vh - 210px);
        }}

        .panel {{
            width: min(760px, 100%);
            position: relative;
            padding: 1px;
            border-radius: 34px;
            background:
                linear-gradient(145deg, rgba(110, 231, 255, 0.55), rgba(139, 92, 246, 0.2), rgba(236, 72, 153, 0.45));
            box-shadow: var(--shadow);
        }}

        .panel-inner {{
            position: relative;
            padding: 42px;
            border-radius: 33px;
            border: 1px solid var(--card-border);
            background:
                linear-gradient(180deg, rgba(255, 255, 255, 0.13), rgba(255, 255, 255, 0.07));
            backdrop-filter: blur(24px);
            overflow: hidden;
            text-align: center;
        }}

        .panel-inner::before {{
            content: "";
            position: absolute;
            inset: -100px -120px auto auto;
            width: 280px;
            height: 280px;
            background: radial-gradient(circle, rgba(110, 231, 255, 0.18), transparent 68%);
            pointer-events: none;
        }}

        .success-icon {{
            position: relative;
            width: 92px;
            height: 92px;
            display: grid;
            place-items: center;
            margin: 0 auto 24px;
            border-radius: 30px;
            background:
                linear-gradient(135deg, rgba(52, 211, 153, 0.95), rgba(110, 231, 255, 0.95));
            box-shadow: 0 22px 55px rgba(52, 211, 153, 0.2);
        }}

        h1 {{
            position: relative;
            margin: 0;
            font-size: clamp(38px, 5vw, 58px);
            line-height: 1;
            letter-spacing: -2.2px;
        }}

        .gradient-text {{
            color: #67e8f9;
            text-shadow: 0 0 22px rgba(103, 232, 249, 0.45);
        }}

        .subtitle {{
            position: relative;
            margin: 18px auto 0;
            max-width: 560px;
            color: var(--muted);
            font-size: 17px;
            line-height: 1.65;
        }}

        .file-card {{
            position: relative;
            margin: 26px auto 0;
            padding: 16px 18px;
            border-radius: 22px;
            background: rgba(0, 0, 0, 0.18);
            border: 1px solid rgba(255, 255, 255, 0.1);
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 12px;
            color: #eef2ff;
            word-break: break-word;
        }}

        .file-name {{
            font-weight: 800;
        }}

        .actions {{
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 14px;
            margin-top: 30px;
            flex-wrap: wrap;
        }}

        .btn {{
            min-height: 56px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            gap: 10px;
            padding: 0 24px;
            border: none;
            border-radius: 18px;
            color: #06101f;
            background: linear-gradient(135deg, #67e8f9, #a78bfa 55%, #f0abfc);
            box-shadow:
                0 18px 46px rgba(139, 92, 246, 0.26),
                inset 0 1px 0 rgba(255, 255, 255, 0.55);
            font-size: 16px;
            font-weight: 900;
            cursor: pointer;
            text-decoration: none;
            transition: 0.22s ease;
            white-space: nowrap;
        }}

        .btn:hover {{
            transform: translateY(-2px);
            filter: saturate(1.08);
            box-shadow:
                0 24px 60px rgba(139, 92, 246, 0.36),
                inset 0 1px 0 rgba(255, 255, 255, 0.62);
        }}

        .btn-secondary {{
            min-height: 56px;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            padding: 0 22px;
            border-radius: 18px;
            border: 1px solid rgba(255, 255, 255, 0.16);
            color: var(--text);
            background: rgba(255, 255, 255, 0.08);
            text-decoration: none;
            font-weight: 800;
            transition: 0.2s ease;
        }}

        .btn-secondary:hover {{
            transform: translateY(-2px);
            background: rgba(255, 255, 255, 0.13);
        }}

        .steps {{
            position: relative;
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 12px;
            margin-top: 30px;
        }}

        .step {{
            padding: 14px;
            border-radius: 18px;
            background: rgba(255, 255, 255, 0.065);
            border: 1px solid rgba(255, 255, 255, 0.1);
            text-align: left;
        }}

        .step strong {{
            display: block;
            margin-bottom: 4px;
            font-size: 14px;
        }}

        .step span {{
            color: var(--muted);
            font-size: 13px;
            line-height: 1.45;
        }}

        @media (max-width: 720px) {{
            .topbar {{
                align-items: flex-start;
                flex-direction: column;
                margin-bottom: 36px;
            }}

            .panel-inner {{
                padding: 28px 20px;
            }}

            .steps {{
                grid-template-columns: 1fr;
            }}

            .btn,
            .btn-secondary {{
                width: 100%;
            }}
        }}
    </style>
</head>

<body>
    <div class="orb one"></div>
    <div class="orb two"></div>
    <div class="orb three"></div>

    <main class="page">
        <header class="topbar">
            <div class="brand">
                <div class="brand-icon">
                    <svg width="24" height="24" viewBox="0 0 24 24" fill="none">
                        <path d="M12 3L20 7V12C20 17 16.5 20.5 12 21C7.5 20.5 4 17 4 12V7L12 3Z" stroke="white" stroke-width="2" stroke-linejoin="round"/>
                        <path d="M9 12L11 14L15.5 9.5" stroke="white" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
                    </svg>
                </div>
                <div>
                    <div>PD Redactor</div>
                    <small style="color: var(--muted); font-weight: 600;">Secure document anonymization</small>
                </div>
            </div>

            <div class="status-pill">
                <span class="pulse"></span>
                Обработка завершена
            </div>
        </header>

        <section class="result-wrap">
            <div class="panel">
                <div class="panel-inner">
                    <div class="success-icon">
                        <svg width="44" height="44" viewBox="0 0 24 24" fill="none">
                            <path d="M20 6L9 17L4 12" stroke="white" stroke-width="2.7" stroke-linecap="round" stroke-linejoin="round"/>
                        </svg>
                    </div>

                    <h1>
                        Файл
                        <span class="gradient-text">готов</span>
                    </h1>

                    <p class="subtitle">
                        Обезличенная версия документа успешно сформирована. Теперь файл можно скачать
                        и использовать для безопасной передачи или хранения.
                    </p>

                    <div class="file-card">
                        <span>📄</span>
                        <span class="file-name">{filename}</span>
                    </div>

                    <div class="actions">
                        <a class="btn" href="/api/download/{token}">
                            <span>⬇</span>
                            Скачать файл
                        </a>

                        <a class="btn-secondary" href="/">
                            Обработать ещё
                        </a>
                    </div>

                    <div class="steps">
                        <div class="step">
                            <strong>1. Detection</strong>
                            <span>Персональные данные были обнаружены системой.</span>
                        </div>

                        <div class="step">
                            <strong>2. Redaction</strong>
                            <span>Фрагменты были скрыты согласно выбранной политике.</span>
                        </div>

                        <div class="step">
                            <strong>3. Delivery</strong>
                            <span>Готовый документ доступен для скачивания.</span>
                        </div>
                    </div>
                </div>
            </div>
        </section>
    </main>
</body>
</html>
"""
    return HTMLResponse(html)


@app.get("/api/download/{token}")
async def api_download(token: str, request: Request):
    """
    Отдаём файл из памяти как attachment.
    Фикс для не-ASCII имён: используем filename* (RFC 5987) + ASCII-фолбэк.
    """
    item = DOWNLOAD_CACHE.get(token)
    if not item:
        raise HTTPException(404, "Файл не найден или ссылка устарела.")

    data: bytes = item["bytes"]  # байты файла
    media: str = item["media"]  # MIME типа "application/pdf" или DOCX
    filename: str = item[
        "filename"
    ]  # имя *.redacted.pdf|docx (может быть на кириллице)

    # ASCII-фолбэк (на случай кириллицы): убираем диакритику/не-ascii
    ascii_fallback = (
        unicodedata.normalize("NFKD", filename)
        .encode("ascii", "ignore")
        .decode("ascii")
    )
    if not ascii_fallback:
        ascii_fallback = "download" + Path(filename).suffix

    # RFC 5987: filename* с URL-экранированием UTF-8
    cd = (
        f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"
    )

    headers = {
        "Content-Disposition": cd,
        "Cache-Control": "no-store",
    }

    return StreamingResponse(BytesIO(data), media_type=media, headers=headers)


# --------------------------- Точка входа -----------------------------------
@app.post("/api/train")
async def api_train(
    rows: int = Form(1000),
    policy: str = Form("mask"),
    languages: str = Form("ru,en"),
):
    from app.pipeline.train import train_from_fake

    langs = [l for l in languages.split(",") if l]  # noqa: E741
    metrics = train_from_fake(n_rows=rows, policy=policy, languages=langs)
    return metrics


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app.main:app", host="127.0.0.1", port=8000, reload=True)
